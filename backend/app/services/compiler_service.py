"""
Compiler service — assembles approved chapters into docx or txt.
Ported from compiler.py. Accepts a storage_backend dependency.
"""

import re
from pathlib import Path

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exception_handlers import ConflictError
from app.db.models import Book, BookStatus, OutputFormat
from app.services.chapter_service import get_chapters

logger = structlog.get_logger(__name__)


async def compile_book(
    db: AsyncSession,
    book: Book,
    output_format: OutputFormat = OutputFormat.DOCX,
    output_dir: str = "/app/output",
) -> str:
    """
    Compile the book into a file.
    Returns the local file path (dev) or S3 key (production).
    In production (STORAGE_BACKEND=s3) the file is uploaded to S3 and
    compiled_path stores the S3 key instead of a local path.
    """
    from app.config import settings

    chapters = await get_chapters(db, book.id)

    if not chapters:
        raise ConflictError("No chapters to compile")

    unapproved = [c for c in chapters if not c.approved]
    if unapproved:
        nums = [str(c.number) for c in unapproved]
        raise ConflictError(f"Chapters not yet approved: {', '.join(nums)}")

    filename = f"{book.id}.{output_format.value}"

    if settings.storage_backend.value == "s3":
        # Write to a temp file then upload to S3
        import tempfile, boto3
        with tempfile.NamedTemporaryFile(suffix=f".{output_format.value}", delete=False) as tmp:
            tmp_path = tmp.name

        if output_format == OutputFormat.DOCX:
            _write_docx(book, chapters, tmp_path)
        else:
            _write_txt(book, chapters, tmp_path)

        s3_key = f"books/{book.user_id}/{filename}"
        s3 = boto3.client("s3", region_name=settings.aws_region)
        import os
        try:
            s3.upload_file(tmp_path, settings.aws_s3_bucket, s3_key)
        finally:
            os.unlink(tmp_path)  # always clean up temp file

        book.compiled_path = s3_key   # S3 key, not local path
        book.output_format = output_format
        book.status = BookStatus.COMPLETE
        db.add(book)

        logger.info("book_compiled_s3", book_id=book.id, s3_key=s3_key)
        return s3_key
    else:
        # Local storage (dev)
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        filepath = str(Path(output_dir) / filename)

        if output_format == OutputFormat.DOCX:
            _write_docx(book, chapters, filepath)
        else:
            _write_txt(book, chapters, filepath)

        book.compiled_path = filepath
        book.output_format = output_format
        book.status = BookStatus.COMPLETE
        db.add(book)

        logger.info("book_compiled_local", book_id=book.id, path=filepath)
        return filepath


def _write_docx(book: Book, chapters: list, filepath: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_before = Pt(200)
    run = title_para.add_run(book.title)
    run.font.size = Pt(28)
    run.bold = True
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for ch in chapters:
        doc.add_paragraph(f"Chapter {ch.number}: {ch.title}", style="List Number")
    doc.add_page_break()

    # Chapters
    for ch in chapters:
        doc.add_heading(f"Chapter {ch.number}: {ch.title}", level=1)
        for para in (ch.content or "").split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        doc.add_page_break()

    doc.save(filepath)


def _write_txt(book: Book, chapters: list, filepath: str) -> None:
    lines = [
        "=" * 60,
        book.title.center(60),
        "=" * 60,
        "",
        "TABLE OF CONTENTS",
        "-" * 40,
    ]
    for ch in chapters:
        lines.append(f"  Chapter {ch.number}: {ch.title}")
    lines += ["", "=" * 60, ""]

    for ch in chapters:
        lines += [
            f"CHAPTER {ch.number}: {ch.title.upper()}",
            "-" * 40,
            "",
            ch.content or "",
            "",
            "=" * 60,
            "",
        ]

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
